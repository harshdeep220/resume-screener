"""
extractor.py — Text extraction from PDF, DOCX, and TXT files.

Dispatches by file extension and returns clean UTF-8 strings.
Handles scanned PDFs gracefully by logging a warning.
"""

import logging
import unicodedata
import re
from pathlib import Path

import fitz  # PyMuPDF
from docx import Document

logger = logging.getLogger(__name__)

# Minimum character count to consider extraction successful
_MIN_TEXT_LENGTH = 50


def clean_text(text: str) -> str:
    """Strip non-printable chars, collapse whitespace, normalise Unicode (NFC).

    Args:
        text: Raw extracted text.

    Returns:
        Cleaned UTF-8 string.
    """
    # Normalise Unicode to NFC form
    text = unicodedata.normalize("NFC", text)

    # Remove non-printable / control characters (keep newlines, tabs, spaces)
    text = re.sub(r"[^\S \n\t]", " ", text)          # unusual whitespace → space
    text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]", "", text)

    # Collapse runs of whitespace (but preserve single newlines for structure)
    text = re.sub(r"[^\S\n]+", " ", text)             # horizontal whitespace → single space
    text = re.sub(r"\n{3,}", "\n\n", text)            # 3+ newlines → double newline
    text = re.sub(r"[ \t]+\n", "\n", text)            # trailing spaces before newline

    return text.strip()


def extract_pdf(path: Path) -> str:
    """Extract text from a PDF file using PyMuPDF.

    Args:
        path: Path to the PDF file.

    Returns:
        Extracted and cleaned text. Empty string if the PDF is scanned/image-only.
    """
    doc = fitz.open(str(path))
    pages = []
    for page in doc:
        pages.append(page.get_text("text"))
    doc.close()
    return "\n".join(pages)


def extract_docx(path: Path) -> str:
    """Extract text from a DOCX file, including table cells.

    Args:
        path: Path to the DOCX file.

    Returns:
        Extracted and cleaned text.
    """
    doc = Document(str(path))
    parts = []

    # Paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    # Table cells (common in modern resumes for skills grids)
    for table in doc.tables:
        for row in table.rows:
            row_texts = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_texts.append(cell_text)
            if row_texts:
                parts.append(" | ".join(row_texts))

    return "\n".join(parts)


def extract_txt(path: Path) -> str:
    """Extract text from a plain text file.

    Args:
        path: Path to the TXT file.

    Returns:
        File contents as a string.
    """
    return path.read_text(encoding="utf-8", errors="replace")


def extract_text(path: str | Path) -> tuple[str, str]:
    """Extract and clean text from a file (PDF, DOCX, or TXT).

    Dispatches to the appropriate extractor based on file extension.
    Checks for scanned PDFs and logs a warning if text extraction fails.

    Args:
        path: Path to the file.

    Returns:
        A tuple of (filename, cleaned_text). If the file is a scanned PDF
        or an unsupported format, cleaned_text will be an empty string.
    """
    path = Path(path)
    filename = path.name
    suffix = path.suffix.lower()

    extractors = {
        ".pdf": extract_pdf,
        ".docx": extract_docx,
        ".txt": extract_txt,
    }

    extractor = extractors.get(suffix)
    if extractor is None:
        logger.warning(
            "⚠ [%s] Unsupported file format '%s' — skipping.", filename, suffix
        )
        return filename, ""

    try:
        raw_text = extractor(path)
    except Exception as e:
        logger.error("✖ [%s] Extraction failed: %s — skipping.", filename, e)
        return filename, ""

    cleaned = clean_text(raw_text)

    # Check for scanned / image-only PDFs
    if len(cleaned) < _MIN_TEXT_LENGTH:
        if suffix == ".pdf":
            logger.warning(
                "⚠ [%s] appears to be a scanned PDF — text extraction failed. Skipping.",
                filename,
            )
        else:
            logger.warning(
                "⚠ [%s] extracted text is too short (%d chars) — may be empty or corrupted.",
                filename,
                len(cleaned),
            )
        return filename, ""

    return filename, cleaned
